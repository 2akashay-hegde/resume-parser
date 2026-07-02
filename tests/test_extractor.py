from pathlib import Path

import fitz
from docx import Document

from app.extractor import extract_text


def test_extract_text_from_docx(tmp_path: Path) -> None:
    doc_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Senior Software Engineer")
    doc.save(doc_path)

    text = extract_text(str(doc_path))

    assert "John Doe" in text
    assert "Senior Software Engineer" in text


def test_extract_text_from_pdf(tmp_path: Path) -> None:
    pdf_path = tmp_path / "resume.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Jane Smith\nData Scientist")
    doc.save(pdf_path)
    doc.close()

    text = extract_text(str(pdf_path))

    assert "Jane Smith" in text
    assert "Data Scientist" in text
